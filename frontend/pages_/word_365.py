import streamlit as st
from docx import Document
import io

def show_word_365():
    st.title("Edit Docx File")

    # File upload
    uploaded_file = st.file_uploader("Upload a DOCX file", type="docx")
    
    if uploaded_file is not None:
        # Load the uploaded file
        document = Document(io.BytesIO(uploaded_file.read()))
        
        st.write("Document content:")
        # Displaying the content of the document
        for para in document.paragraphs:
            st.write(para.text)
        
        # Modify the document (Example: Add new text)
        new_text = st.text_area("Add new text to the document:")
        
        if st.button("Add Text"):
            # Modify the document with new text
            document.add_paragraph(new_text)
            st.write("Text added!")
        
        # Allow user to download the modified file
        modified_file = io.BytesIO()
        document.save(modified_file)
        modified_file.seek(0)

        st.download_button(
            label="Download Modified Document",
            data=modified_file,
            file_name="modified_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
show_word_365()